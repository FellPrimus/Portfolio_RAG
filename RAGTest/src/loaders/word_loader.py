"""
Word 문서 로더

이 모듈은 Word 파일(.docx)을 로드하고 텍스트를 추출합니다.
LangChain의 문서 형식과 호환되도록 구현되었습니다.
"""

import os
from docx import Document as DocxDocument
from typing import List, Optional
from langchain_core.documents import Document  # LangChain 1.0+


class WordLoader:
    """
    Word 파일(.docx)에서 텍스트를 추출하는 문서 로더

    주요 기능:
    - .docx 파일 읽기
    - 단락, 표, 헤더/푸터 텍스트 추출
    - 메타데이터 추출 (파일명, 단락 수 등)
    - LangChain Document 형식으로 반환
    """

    def __init__(
        self,
        extract_tables: bool = True,
        extract_headers: bool = False
    ):
        """
        WordLoader 초기화

        Args:
            extract_tables (bool): 표 내용 추출 여부. 기본값 True
            extract_headers (bool): 헤더/푸터 추출 여부. 기본값 False
        """
        self.extract_tables = extract_tables
        self.extract_headers = extract_headers

    def load(self, file_path: str) -> List[Document]:
        """
        Word 파일을 로드합니다.

        Args:
            file_path (str): Word 파일 경로

        Returns:
            List[Document]: LangChain Document 객체 리스트

        Raises:
            Exception: Word 로드 실패 시

        Example:
            >>> loader = WordLoader()
            >>> docs = loader.load("./data/sample.docx")
            >>> print(f"{len(docs)}개 문서 로드됨")
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")

            # Word 문서 열기
            doc = DocxDocument(file_path)

            # 텍스트 추출
            text_content = self._extract_text(doc)

            # 메타데이터 생성
            metadata = self._create_metadata(file_path, doc)

            # Document 생성
            document = Document(
                page_content=text_content,
                metadata=metadata
            )

            return [document]

        except Exception as e:
            raise Exception(f"Word 로드 실패: {file_path}\n에러: {str(e)}")

    def load_multiple(self, file_paths: List[str]) -> List[Document]:
        """
        여러 Word 파일을 로드합니다.

        Args:
            file_paths (List[str]): Word 파일 경로 리스트

        Returns:
            List[Document]: 모든 파일에서 로드된 Document 리스트

        Example:
            >>> loader = WordLoader()
            >>> files = ["./data/doc1.docx", "./data/doc2.docx"]
            >>> docs = loader.load_multiple(files)
            >>> print(f"총 {len(docs)}개 문서 로드됨")
        """
        all_documents = []

        for file_path in file_paths:
            try:
                documents = self.load(file_path)
                all_documents.extend(documents)
                print(f"✓ 로드 완료: {file_path}")
            except Exception as e:
                print(f"✗ 로드 실패: {file_path} - {str(e)}")
                continue

        return all_documents

    def _extract_text(self, doc: DocxDocument) -> str:
        """
        Word 문서에서 텍스트를 추출합니다.

        Args:
            doc (DocxDocument): python-docx Document 객체

        Returns:
            str: 추출된 텍스트
        """
        text_parts = []

        # 1. 단락(paragraphs) 추출
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)

        # 2. 표(tables) 추출
        if self.extract_tables:
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    text_parts.append(table_text)

        # 3. 헤더/푸터 추출 (옵션)
        if self.extract_headers:
            for section in doc.sections:
                # 헤더
                header_text = self._extract_header_footer_text(section.header)
                if header_text:
                    text_parts.append(f"[Header] {header_text}")

                # 푸터
                footer_text = self._extract_header_footer_text(section.footer)
                if footer_text:
                    text_parts.append(f"[Footer] {footer_text}")

        return '\n\n'.join(text_parts)

    def _extract_table_text(self, table) -> str:
        """
        표에서 텍스트를 추출합니다.

        Args:
            table: python-docx Table 객체

        Returns:
            str: 표 내용을 텍스트로 변환
        """
        table_data = []

        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)

            if any(row_data):  # 빈 행이 아닌 경우만
                table_data.append(" | ".join(row_data))

        return '\n'.join(table_data) if table_data else ""

    def _extract_header_footer_text(self, header_or_footer) -> str:
        """
        헤더 또는 푸터에서 텍스트를 추출합니다.

        Args:
            header_or_footer: python-docx Header 또는 Footer 객체

        Returns:
            str: 추출된 텍스트
        """
        text_parts = []

        for paragraph in header_or_footer.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)

        return ' '.join(text_parts)

    def _create_metadata(self, file_path: str, doc: DocxDocument) -> dict:
        """
        Word 메타데이터를 생성합니다.

        Args:
            file_path (str): 파일 경로
            doc (DocxDocument): python-docx Document 객체

        Returns:
            dict: 메타데이터 딕셔너리
        """
        metadata = {
            'source': file_path,
            'file_name': os.path.basename(file_path),
            'type': 'word',
            'num_paragraphs': len(doc.paragraphs),
            'num_tables': len(doc.tables)
        }

        # Core properties 추출 (생성자, 제목 등)
        try:
            core_props = doc.core_properties
            if core_props.title:
                metadata['title'] = core_props.title
            if core_props.author:
                metadata['author'] = core_props.author
            if core_props.created:
                metadata['created'] = str(core_props.created)
            if core_props.modified:
                metadata['modified'] = str(core_props.modified)
        except:
            pass  # 메타데이터가 없을 수 있음

        return metadata


# 사용 예제
if __name__ == "__main__":
    print("=== Word 문서 로더 예제 ===\n")

    loader = WordLoader(extract_tables=True)

    try:
        docs = loader.load("./data/sample.docx")
        print(f"문서 로드 성공!")
        print(f"파일명: {docs[0].metadata.get('file_name', 'N/A')}")
        print(f"단락 수: {docs[0].metadata.get('num_paragraphs', 'N/A')}")
        print(f"표 수: {docs[0].metadata.get('num_tables', 'N/A')}")
        print(f"제목: {docs[0].metadata.get('title', 'N/A')}")
        print(f"내용 미리보기: {docs[0].page_content[:200]}...")
    except Exception as e:
        print(f"에러: {e}")
