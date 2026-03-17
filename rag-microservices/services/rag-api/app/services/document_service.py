"""
Document Service - Document Management

Handles document upload, chunking, and indexing
"""

import logging
import uuid
import hashlib
from datetime import datetime
from typing import Dict, List, Optional, Any, BinaryIO

from ..utils.clients import EmbeddingClient, QdrantService, RedisService

logger = logging.getLogger(__name__)

# 한글 띄어쓰기 보정기 (싱글톤)
_kiwi_instance = None

def get_kiwi():
    """Kiwi 인스턴스 반환 (지연 로딩)"""
    global _kiwi_instance
    if _kiwi_instance is None:
        try:
            from kiwipiepy import Kiwi
            _kiwi_instance = Kiwi()
            logger.info("한글 띄어쓰기 보정기 (kiwipiepy) 초기화 완료")
        except Exception as e:
            logger.warning(f"한글 띄어쓰기 보정기 초기화 실패: {e}")
            _kiwi_instance = False  # 실패 시 재시도 방지
    return _kiwi_instance if _kiwi_instance else None


def correct_korean_spacing(text: str) -> str:
    """한글 텍스트 띄어쓰기 보정 (kiwipiepy 사용)"""
    if not text:
        return text

    kiwi = get_kiwi()
    if not kiwi:
        return text

    try:
        # 문단 단위로 처리
        paragraphs = text.split('\n')
        corrected_paragraphs = []

        for para in paragraphs:
            if para.strip():
                # 한글이 포함된 경우에만 보정
                if any('\uac00' <= char <= '\ud7a3' for char in para):
                    try:
                        # kiwipiepy의 space 메서드로 띄어쓰기 보정
                        corrected = kiwi.space(para, reset_whitespace=False)
                        corrected_paragraphs.append(corrected)
                    except Exception:
                        corrected_paragraphs.append(para)
                else:
                    corrected_paragraphs.append(para)
            else:
                corrected_paragraphs.append(para)

        return '\n'.join(corrected_paragraphs)
    except Exception as e:
        logger.warning(f"띄어쓰기 보정 실패: {e}")
        return text


class TextChunker:
    """Simple text chunker"""

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        separator: str = "\n"
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separator = separator

    def split_text(self, text: str) -> List[str]:
        """Split text into chunks"""
        if not text:
            return []

        # Split by separator first
        paragraphs = text.split(self.separator)
        chunks = []
        current_chunk = []
        current_length = 0

        for para in paragraphs:
            para_length = len(para)

            if current_length + para_length > self.chunk_size and current_chunk:
                # Save current chunk
                chunks.append(self.separator.join(current_chunk))

                # Start new chunk with overlap
                overlap_text = self.separator.join(current_chunk)[-self.chunk_overlap:]
                current_chunk = [overlap_text] if overlap_text else []
                current_length = len(overlap_text) if overlap_text else 0

            current_chunk.append(para)
            current_length += para_length + len(self.separator)

        # Add remaining
        if current_chunk:
            chunks.append(self.separator.join(current_chunk))

        return chunks


class DocumentService:
    """
    Document Management Service

    Features:
    - Document upload and parsing
    - Text chunking
    - Embedding generation
    - Vector storage in Qdrant
    - Metadata storage in Redis
    """

    def __init__(
        self,
        embedding_client: EmbeddingClient,
        qdrant_service: QdrantService,
        redis_service: RedisService,
        chunk_size: int = 1000,
        chunk_overlap: int = 200
    ):
        self.embedding_client = embedding_client
        self.qdrant_service = qdrant_service
        self.redis_service = redis_service
        self.chunker = TextChunker(chunk_size, chunk_overlap)

    def _generate_doc_id(self, content: str) -> str:
        """Generate unique document ID"""
        return str(uuid.uuid4())

    def _generate_chunk_id(self, doc_id: str, chunk_index: int) -> str:
        """Generate unique chunk ID"""
        return f"{doc_id}_chunk_{chunk_index}"

    async def _extract_text(self, file: BinaryIO, filename: str) -> str:
        """Extract text from file"""
        import re
        extension = filename.lower().split(".")[-1]

        if extension in ["txt", "md"]:
            content = file.read()
            if isinstance(content, bytes):
                content = content.decode("utf-8", errors="ignore")
            return content

        elif extension in ["htm", "html"]:
            content = file.read()
            if isinstance(content, bytes):
                content = content.decode("utf-8", errors="ignore")
            # Simple HTML tag removal
            content = re.sub(r"<[^>]+>", " ", content)
            content = re.sub(r"\s+", " ", content)
            return content.strip()

        elif extension == "pdf":
            return await self._extract_pdf(file)

        elif extension in ["xlsx", "xls"]:
            return await self._extract_excel(file, extension)

        elif extension in ["docx", "doc"]:
            return await self._extract_word(file, extension)

        else:
            # Try to read as text
            content = file.read()
            if isinstance(content, bytes):
                content = content.decode("utf-8", errors="ignore")
            return content

    async def _extract_pdf(self, file: BinaryIO) -> str:
        """Extract text from PDF file"""
        try:
            from pypdf import PdfReader
            import io

            # Ensure we have bytes
            content = file.read()
            if not isinstance(content, bytes):
                content = content.encode()

            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)

            text_parts = []
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        # 한글 띄어쓰기 보정 적용
                        page_text = correct_korean_spacing(page_text)
                        text_parts.append(f"[페이지 {page_num + 1}]\n{page_text}")
                except Exception as e:
                    logger.warning(f"PDF 페이지 {page_num + 1} 추출 실패: {e}")

            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"PDF 텍스트 추출 실패: {e}")
            return ""

    async def _extract_excel(self, file: BinaryIO, extension: str) -> str:
        """Extract text from Excel file"""
        try:
            import io

            content = file.read()
            if not isinstance(content, bytes):
                content = content.encode()

            excel_file = io.BytesIO(content)

            if extension == "xlsx":
                from openpyxl import load_workbook
                wb = load_workbook(excel_file, read_only=True, data_only=True)
                text_parts = []

                for sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]
                    sheet_text = [f"[시트: {sheet_name}]"]

                    for row in sheet.iter_rows(values_only=True):
                        row_values = [str(cell) if cell is not None else "" for cell in row]
                        if any(row_values):
                            sheet_text.append("\t".join(row_values))

                    text_parts.append("\n".join(sheet_text))
                wb.close()
                return "\n\n".join(text_parts)

            else:  # xls
                import xlrd
                wb = xlrd.open_workbook(file_contents=content)
                text_parts = []

                for sheet_idx in range(wb.nsheets):
                    sheet = wb.sheet_by_index(sheet_idx)
                    sheet_text = [f"[시트: {sheet.name}]"]

                    for row_idx in range(sheet.nrows):
                        row_values = [str(sheet.cell_value(row_idx, col_idx)) for col_idx in range(sheet.ncols)]
                        if any(row_values):
                            sheet_text.append("\t".join(row_values))

                    text_parts.append("\n".join(sheet_text))
                return "\n\n".join(text_parts)

        except Exception as e:
            logger.error(f"Excel 텍스트 추출 실패: {e}")
            return ""

    async def _extract_word(self, file: BinaryIO, extension: str) -> str:
        """Extract text from Word file"""
        try:
            import io

            content = file.read()
            if not isinstance(content, bytes):
                content = content.encode()

            if extension == "docx":
                from docx import Document
                doc_file = io.BytesIO(content)
                doc = Document(doc_file)

                text_parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)

                # 테이블 내용도 추출
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells]
                        if any(row_text):
                            text_parts.append("\t".join(row_text))

                return "\n".join(text_parts)

            else:  # doc (구 형식)
                # doc 파일은 python-docx로 직접 처리 불가
                # 간단한 텍스트 추출 시도
                logger.warning("구 Word 형식(.doc)은 제한적으로 지원됩니다. .docx 형식 사용을 권장합니다.")
                try:
                    # 바이너리에서 텍스트 추출 시도
                    text = content.decode("utf-8", errors="ignore")
                    import re
                    # 특수 문자 제거
                    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]', ' ', text)
                    text = re.sub(r'\s+', ' ', text)
                    return text.strip()
                except:
                    return ""

        except Exception as e:
            logger.error(f"Word 텍스트 추출 실패: {e}")
            return ""

    async def upload_document(
        self,
        file: BinaryIO,
        filename: str,
        category: str = "general",
        metadata: Dict[str, Any] = None,
        collection: str = "documents"
    ) -> Dict[str, Any]:
        """
        Upload and index a document

        Args:
            file: File object
            filename: Original filename
            category: Document category
            metadata: Additional metadata
            collection: Qdrant collection name

        Returns:
            Dict with document_id, status, chunks_count
        """
        # Extract text
        logger.info(f"Extracting text from: {filename}")
        text = await self._extract_text(file, filename)

        if not text or len(text.strip()) < 10:
            return {
                "document_id": None,
                "status": "error",
                "chunks_count": 0,
                "message": "Could not extract text from document"
            }

        # Generate document ID
        doc_id = self._generate_doc_id(text)

        # Chunk text
        chunks = self.chunker.split_text(text)
        logger.info(f"Created {len(chunks)} chunks for document {doc_id}")

        if not chunks:
            return {
                "document_id": doc_id,
                "status": "error",
                "chunks_count": 0,
                "message": "No chunks created from document"
            }

        # Generate embeddings
        logger.info(f"Generating embeddings for {len(chunks)} chunks...")
        embeddings = await self.embedding_client.embed_documents(chunks)

        # Ensure collection exists
        self.qdrant_service.ensure_collection(collection)

        # Prepare vectors with metadata
        chunk_ids = []
        payloads = []
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            chunk_id = self._generate_chunk_id(doc_id, i)
            chunk_ids.append(chunk_id)
            payloads.append({
                "content": chunk,
                "doc_id": doc_id,
                "chunk_index": i,
                "source": filename,
                "category": category,
                "created_at": datetime.now().isoformat(),
                **(metadata or {})
            })

        # Store in Qdrant
        logger.info(f"Storing {len(chunks)} vectors in Qdrant...")
        self.qdrant_service.upsert_vectors(
            ids=chunk_ids,
            vectors=embeddings,
            payloads=payloads,
            collection_name=collection
        )

        # Store metadata in Redis
        doc_meta = {
            "id": doc_id,
            "filename": filename,
            "category": category,
            "chunk_count": str(len(chunks)),
            "created_at": datetime.now().isoformat(),
            "collection": collection,
            **(metadata or {})
        }
        self.redis_service.set_document_meta(doc_id, doc_meta)

        # Update category count
        self.redis_service.increment_category_count(category)

        return {
            "document_id": doc_id,
            "status": "success",
            "chunks_count": len(chunks),
            "message": f"Document indexed successfully with {len(chunks)} chunks"
        }

    async def upload_text(
        self,
        text: str,
        source: str,
        category: str = "general",
        metadata: Dict[str, Any] = None,
        collection: str = "documents"
    ) -> Dict[str, Any]:
        """
        Upload and index plain text

        Args:
            text: Text content
            source: Source identifier
            category: Document category
            metadata: Additional metadata
            collection: Qdrant collection name
        """
        if not text or len(text.strip()) < 10:
            return {
                "document_id": None,
                "status": "error",
                "chunks_count": 0,
                "message": "Text too short"
            }

        doc_id = self._generate_doc_id(text)
        chunks = self.chunker.split_text(text)

        if not chunks:
            return {
                "document_id": doc_id,
                "status": "error",
                "chunks_count": 0,
                "message": "No chunks created"
            }

        # Generate embeddings
        embeddings = await self.embedding_client.embed_documents(chunks)

        # Ensure collection exists
        self.qdrant_service.ensure_collection(collection)

        # Prepare vectors
        chunk_ids = []
        payloads = []
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            chunk_id = self._generate_chunk_id(doc_id, i)
            chunk_ids.append(chunk_id)
            payloads.append({
                "content": chunk,
                "doc_id": doc_id,
                "chunk_index": i,
                "source": source,
                "category": category,
                "created_at": datetime.now().isoformat(),
                **(metadata or {})
            })

        # Store in Qdrant
        self.qdrant_service.upsert_vectors(
            ids=chunk_ids,
            vectors=embeddings,
            payloads=payloads,
            collection_name=collection
        )

        # Store metadata
        doc_meta = {
            "id": doc_id,
            "filename": source,
            "category": category,
            "chunk_count": str(len(chunks)),
            "created_at": datetime.now().isoformat(),
            "collection": collection
        }
        self.redis_service.set_document_meta(doc_id, doc_meta)
        self.redis_service.increment_category_count(category)

        return {
            "document_id": doc_id,
            "status": "success",
            "chunks_count": len(chunks),
            "message": f"Text indexed with {len(chunks)} chunks"
        }

    def get_document(self, doc_id: str) -> Optional[Dict[str, Any]]:
        """Get document metadata"""
        return self.redis_service.get_document_meta(doc_id)

    def list_documents(
        self,
        category: str = None,
        page: int = 1,
        page_size: int = 20
    ) -> Dict[str, Any]:
        """List documents with pagination"""
        # Get all document keys
        all_docs = []
        for key in self.redis_service.client.keys("doc:*"):
            doc = self.redis_service.client.hgetall(key)
            if doc:
                if category is None or doc.get("category") == category:
                    all_docs.append(doc)

        # Sort by created_at descending
        all_docs.sort(key=lambda x: x.get("created_at", ""), reverse=True)

        # Paginate
        start = (page - 1) * page_size
        end = start + page_size
        paginated = all_docs[start:end]

        return {
            "documents": paginated,
            "total": len(all_docs),
            "page": page,
            "page_size": page_size
        }

    def delete_document(
        self,
        doc_id: str,
        collection: str = "documents"
    ) -> Dict[str, Any]:
        """Delete document and its chunks"""
        # Get document metadata
        doc_meta = self.redis_service.get_document_meta(doc_id)
        if not doc_meta:
            return {
                "success": False,
                "deleted_chunks": 0,
                "message": "Document not found"
            }

        # Delete from Qdrant
        deleted_count = self.qdrant_service.delete_by_filter(
            filter_key="doc_id",
            filter_value=doc_id,
            collection_name=collection
        )

        # Delete metadata from Redis
        self.redis_service.delete_document_meta(doc_id)

        # Update category count
        category = doc_meta.get("category")
        if category:
            chunk_count = int(doc_meta.get("chunk_count", 0))
            self.redis_service.increment_category_count(category, -chunk_count)

        return {
            "success": True,
            "deleted_chunks": deleted_count,
            "message": f"Deleted document {doc_id} with {deleted_count} chunks"
        }
